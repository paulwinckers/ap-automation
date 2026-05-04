"""
Handoff Pack Generator
======================
Fetches data from Aspire (Opportunities, Services, WorkTickets, Receipts)
and produces a professional .docx Project Handoff Package for download.

Route: GET /handoff/generate?opportunity_number=1970
Returns: application/vnd.openxmlformats-officedocument.wordprocessingml.document
"""
import io
import json
import logging
from datetime import date
from typing import Any, Dict, List, Optional

import anthropic as _anthropic
from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import StreamingResponse

from app.core.config import settings
from app.services.aspire import AspireClient

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/handoff", tags=["handoff"])

_aspire = AspireClient(sandbox=settings.ASPIRE_DASHBOARD_SANDBOX)

# ── AI client (lazy singleton) ────────────────────────────────────────────────

_ai_client: Optional[_anthropic.AsyncAnthropic] = None

def _get_ai() -> _anthropic.AsyncAnthropic:
    global _ai_client
    if _ai_client is None:
        _ai_client = _anthropic.AsyncAnthropic(api_key=settings.ANTHROPIC_API_KEY)
    return _ai_client


async def _generate_scope_summary(
    opp_name:     str,
    prop_name:    str,
    proposal_desc: str,
    group_names:  List[str],
    total_hours:  float,
) -> Dict[str, str]:
    """
    Ask Claude to produce a concise project summary from the Aspire data.
    Returns dict with: project_description, key_objectives, major_deliverables, exclusions
    """
    context_parts = [f"Project: {opp_name}"]
    if prop_name:
        context_parts.append(f"Property: {prop_name}")
    if proposal_desc:
        context_parts.append(f"Proposal description: {proposal_desc[:800]}")
    if group_names:
        context_parts.append(f"Scope phases/groups: {', '.join(group_names)}")
    if total_hours:
        context_parts.append(f"Total estimated hours: {total_hours:,.1f}")
    context = "\n".join(context_parts)

    prompt = (
        "You are a project manager summarizing a landscaping project handoff document.\n"
        "Based on the following Aspire project data, produce a concise project overview.\n\n"
        f"{context}\n\n"
        "Respond ONLY with valid JSON — no markdown, no explanation:\n"
        "{\n"
        '  "project_description": "One sentence describing the overall project",\n'
        '  "key_objectives": "Comma-separated list of the main work areas / deliverables",\n'
        '  "major_deliverables": "Key notes about the job — billing type, documentation needs, etc.",\n'
        '  "exclusions": "Any obvious exclusions or leave blank"\n'
        "}"
    )

    try:
        msg = await _get_ai().messages.create(
            model="claude-opus-4-5",
            max_tokens=400,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        if raw.startswith("```"):
            parts = raw.split("```")
            raw = parts[1].lstrip("json").strip() if len(parts) > 1 else raw
        return json.loads(raw)
    except Exception as exc:
        logger.warning(f"AI scope summary failed: {exc}")
        return {
            "project_description": opp_name,
            "key_objectives":      ", ".join(group_names) if group_names else "—",
            "major_deliverables":  "",
            "exclusions":          "",
        }

# ── Brand colours (hex without #) ────────────────────────────────────────────
NAVY   = "1B3A5C"
LIGHT  = "EBF0F8"
WHITE  = "FFFFFF"
GREEN  = "D4EDDA"
AMBER  = "FFF3CD"
GREY   = "F8F9FA"
MID    = "DEE2E6"

# ── Aspire data helpers ───────────────────────────────────────────────────────

def _check_credentials():
    if not settings.ASPIRE_CLIENT_ID or not settings.ASPIRE_CLIENT_SECRET:
        raise HTTPException(status_code=503, detail="Aspire credentials not configured")


def _str(v: Any, fallback: str = "—") -> str:
    if v is None or v == "":
        return fallback
    return str(v)


def _strip_html(html_text: Any) -> str:
    """Strip HTML tags and decode entities from Aspire rich-text fields."""
    if not html_text:
        return ""
    import html as _html
    import re
    text = str(html_text)
    # Replace block-level tags with newlines so structure is preserved
    text = re.sub(r"<(?:br|p|div|h[1-6]|blockquote|li)[^>]*>", "\n", text, flags=re.IGNORECASE)
    # Strip all remaining tags
    text = re.sub(r"<[^>]+>", "", text)
    # Decode HTML entities (&nbsp; &amp; etc.)
    text = _html.unescape(text)
    # Collapse multiple blank lines and strip leading/trailing whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _fmt_date(d: Optional[str]) -> str:
    if not d:
        return "—"
    try:
        # Aspire often returns ISO with time — truncate
        return d[:10]
    except Exception:
        return d


def _fmt_money(v: Any) -> str:
    if v is None:
        return "—"
    try:
        return f"${float(v):,.2f}"
    except Exception:
        return str(v)


async def _fetch_catalog_vendors(catalog_ids: List[int]) -> Dict[int, str]:
    """
    Given a list of CatalogItemIDs, return a map of CatalogItemID → VendorName.
    Fetches in one OData call; falls back to empty map on error.
    """
    if not catalog_ids:
        return {}
    unique_ids = list(set(catalog_ids))[:50]  # cap to avoid URL length issues
    id_list = ",".join(str(i) for i in unique_ids)
    try:
        res = await _aspire._get("CatalogItems", {
            "$filter": f"CatalogItemID in ({id_list})",
            "$top": "100",
        })
        rows = _aspire._extract_list(res)
        if rows:
            logger.info(f"CatalogItems ALL keys: {sorted(rows[0].keys())}")
            # Log first row values for vendor-related fields to help identify the right one
            vendor_clues = {k: v for k, v in rows[0].items() if any(
                kw in k.lower() for kw in ("vendor", "supplier", "manufacturer", "source", "purchase")
            )}
            logger.info(f"CatalogItems vendor-related fields: {vendor_clues}")
        return {
            r.get("CatalogItemID"): (r.get("VendorName") or r.get("SupplierName") or "")
            for r in rows
            if r.get("CatalogItemID")
        }
    except Exception as e:
        logger.warning(f"CatalogItems vendor fetch failed: {e}")
        return {}


async def _fetch_opportunity(opp_number: int) -> Dict[str, Any]:
    """Fetch opportunity by display number (OpportunityNumber)."""
    res = await _aspire._get("Opportunities", {
        "$filter": f"OpportunityNumber eq {opp_number}",
        "$top": "1",
    })
    items = _aspire._extract_list(res)
    if items:
        return items[0]
    # Fallback: search by OpportunityID
    res2 = await _aspire._get("Opportunities", {
        "$filter": f"OpportunityID eq {opp_number}",
        "$top": "1",
    })
    items2 = _aspire._extract_list(res2)
    if items2:
        return items2[0]
    raise HTTPException(status_code=404, detail=f"Opportunity #{opp_number} not found in Aspire")


async def _fetch_services(opp_id: int) -> List[Dict[str, Any]]:
    try:
        res = await _aspire._get("OpportunityServices", {
            "$filter": f"OpportunityID eq {opp_id}",
            "$orderby": "SortOrder asc",
            "$top": "100",
        })
        return _aspire._extract_list(res)
    except Exception as e:
        logger.warning(f"Could not fetch OpportunityServices: {e}")
        return []


async def _fetch_service_groups(opp_id: int) -> List[Dict[str, Any]]:
    try:
        res = await _aspire._get("OpportunityServiceGroups", {
            "$filter": f"OpportunityID eq {opp_id}",
            "$top": "100",
        })
        rows = _aspire._extract_list(res)
        if rows:
            logger.info(f"OpportunityServiceGroups sample keys: {sorted(rows[0].keys())}")
        return rows
    except Exception as e:
        logger.warning(f"Could not fetch OpportunityServiceGroups: {e}")
        return []


async def _fetch_service_items(opp_id: int, service_ids: List[int]) -> List[Dict[str, Any]]:
    """
    Fetch material/catalog items for all services.
    Tries filtering by OpportunityID first; if empty, falls back to per-service fetches.
    """
    # Attempt 1 — filter by OpportunityID
    for endpoint in ("OpportunityServiceItems", "ServiceItems"):
        try:
            res = await _aspire._get(endpoint, {
                "$filter": f"OpportunityID eq {opp_id}",
                "$top": "500",
            })
            rows = _aspire._extract_list(res)
            if rows:
                logger.info(f"{endpoint} (by opp) sample keys: {sorted(rows[0].keys())}")
                return rows
        except Exception as e:
            logger.debug(f"{endpoint} by opp failed: {e}")

    # Attempt 2 — per-service fetch (Aspire sometimes requires this)
    if not service_ids:
        return []
    all_items: List[Dict[str, Any]] = []
    logged = False
    for sid in service_ids[:20]:  # cap at 20 services to avoid timeout
        for endpoint in ("OpportunityServiceItems", "ServiceItems"):
            try:
                res = await _aspire._get(endpoint, {
                    "$filter": f"OpportunityServiceID eq {sid}",
                    "$top": "100",
                })
                rows = _aspire._extract_list(res)
                if rows:
                    if not logged:
                        logger.info(f"{endpoint} (per-svc) sample keys: {sorted(rows[0].keys())}")
                        logged = True
                    all_items.extend(rows)
                break  # found working endpoint for this service
            except Exception:
                continue
    return all_items


async def _fetch_work_tickets(opp_id: int) -> List[Dict[str, Any]]:
    try:
        res = await _aspire._get("WorkTickets", {
            "$filter": f"OpportunityID eq {opp_id}",
            "$orderby": "ScheduledStartDate asc",
            "$top": "200",
        })
        rows = _aspire._extract_list(res)
        if rows:
            logger.info(f"WorkTickets ALL keys: {sorted(rows[0].keys())}")
            hours_clues = {k: v for k, v in rows[0].items() if any(
                kw in k.lower() for kw in ("hour", "man", "labor", "actual", "estimated", "budget", "scheduled")
            )}
            logger.info(f"WorkTickets hours-related fields: {hours_clues}")
        return rows
    except Exception as e:
        logger.warning(f"Could not fetch WorkTickets: {e}")
        return []


async def _fetch_receipts(opp_id: int) -> List[Dict[str, Any]]:
    """
    Aspire stores purchase receipts linked to opportunities via allocations.
    Try several endpoint + filter combinations until one returns data.
    """
    attempts = [
        # Most likely: PurchaseReceipts filtered by OpportunityID
        ("PurchaseReceipts", {"$filter": f"OpportunityID eq {opp_id}", "$top": "200"}),
        # Via receipt allocations
        ("ReceiptAllocations", {"$filter": f"OpportunityID eq {opp_id}", "$top": "200"}),
        # Legacy endpoint name
        ("Receipts", {"$filter": f"OpportunityID eq {opp_id}", "$top": "200"}),
        # Maybe the field is AllocationOpportunityID
        ("PurchaseReceipts", {"$filter": f"AllocationOpportunityID eq {opp_id}", "$top": "200"}),
    ]

    for endpoint, params in attempts:
        try:
            res  = await _aspire._get(endpoint, params)
            rows = _aspire._extract_list(res)
            if rows:
                logger.info(f"Receipts found via {endpoint} ({params.get('$filter')}): {len(rows)} rows")
                logger.info(f"Receipt ALL keys: {sorted(rows[0].keys())}")
                return rows
            else:
                logger.debug(f"Receipts: {endpoint} returned 0 rows for filter {params.get('$filter')}")
        except Exception as e:
            logger.debug(f"Receipts: {endpoint} failed — {e}")

    logger.warning(f"Could not fetch receipts for opportunity {opp_id} — all attempts returned empty")
    return []


# ── DOCX builder ──────────────────────────────────────────────────────────────

def _build_docx(
    opp:             Dict[str, Any],
    services:        List[Dict[str, Any]],
    service_groups:  List[Dict[str, Any]],
    service_items:   List[Dict[str, Any]],
    catalog_vendors: Dict[int, str],
    scope_summary:   Dict[str, str],
    tickets:         List[Dict[str, Any]],
    receipts:        List[Dict[str, Any]],
) -> bytes:
    """Build the complete handoff pack and return as bytes."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = Document()

    # ── Page setup (US Letter, 1" margins) ────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _rgb(hex_str: str) -> RGBColor:
        h = hex_str.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def _set_cell_bg(cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _set_col_width(table, col_idx: int, width_inches: float):
        """Set all cells in a column to a specific width."""
        from docx.shared import Inches as _Inches
        for row in table.rows:
            row.cells[col_idx].width = _Inches(width_inches)

    def _heading(text: str, level: int = 1):
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = _rgb(NAVY)
        if level == 1:
            run.font.size = Pt(16)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after  = Pt(6)
        else:
            run.font.size = Pt(13)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
        # Navy bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "8" if level == 1 else "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), NAVY)
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def _kv_table(rows: List[tuple]):
        """Two-column key-value table with navy header column."""
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (key, val) in enumerate(rows):
            kc = t.rows[i].cells[0]
            vc = t.rows[i].cells[1]
            kc.width = Inches(2.0)
            vc.width = Inches(4.5)
            _set_cell_bg(kc, LIGHT)
            kp = kc.paragraphs[0]
            kr = kp.add_run(str(key))
            kr.bold = True
            kr.font.size = Pt(10)
            kr.font.color.rgb = _rgb(NAVY)
            vp = vc.paragraphs[0]
            vp.add_run(_str(val))
            vp.runs[0].font.size = Pt(10)
        return t

    def _section_table_header(table, headers: List[str], bg: str = NAVY):
        """Style the first row as a header row."""
        hrow = table.rows[0]
        for i, h in enumerate(headers):
            cell = hrow.cells[i]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = _rgb(WHITE)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Title page ────────────────────────────────────────────────────────────
    opp_name   = _str(opp.get("OpportunityName"), "Project")
    opp_num    = _str(opp.get("OpportunityNumber"), "")
    prop_name  = _str(opp.get("PropertyName"), "")
    today_str  = date.today().strftime("%B %d, %Y")

    # Spacer
    for _ in range(5):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("PROJECT HANDOFF PACKAGE")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = _rgb(NAVY)

    doc.add_paragraph()

    proj_p = doc.add_paragraph()
    proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = proj_p.add_run(opp_name)
    pr.font.size = Pt(20)
    pr.bold = True

    if prop_name and prop_name != "—":
        pp2 = doc.add_paragraph()
        pp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp2.add_run(prop_name).font.size = Pt(14)

    if opp_num and opp_num != "—":
        pnum = doc.add_paragraph()
        pnum.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = pnum.add_run(f"Opportunity #{opp_num}")
        nr.font.size = Pt(12)
        nr.font.color.rgb = _rgb("6c757d")

    for _ in range(4):
        doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(f"Prepared: {today_str}")
    dr.font.size = Pt(11)
    dr.font.color.rgb = _rgb("6c757d")

    doc.add_page_break()

    # ── Section 1: Project Overview ───────────────────────────────────────────
    _heading("1. Project Overview")

    overview_rows = [
        ("Project Name",         opp.get("OpportunityName")),
        ("Property",             opp.get("PropertyName")),
        ("Opportunity #",        opp.get("OpportunityNumber")),
        ("Status",               opp.get("OpportunityStatusName")),
        ("Job Status",           opp.get("JobStatusName")),
        ("Division",             opp.get("DivisionName") or opp.get("BranchName")),
        ("Contract Value",       _fmt_money(opp.get("WonDollars") or opp.get("EstimatedDollars"))),
        ("Start Date",           _fmt_date(opp.get("StartDate"))),
        ("End Date",             _fmt_date(opp.get("EndDate"))),
        ("Won Date",             _fmt_date(opp.get("WonDate"))),
        ("Sales Rep",            opp.get("SalesRepContactName")),
        ("Operations Manager",   opp.get("OperationsManagerContactName")),
        ("Billing Contact",      opp.get("BillingContactName")),
        ("Project Manager",      opp.get("ProjectManagerContactName")),
    ]
    _kv_table(overview_rows)
    doc.add_paragraph()

    # ── Section 2: Scope of Work ──────────────────────────────────────────────
    doc.add_page_break()
    _heading("2. Scope of Work")

    # ── AI-generated project overview box ─────────────────────────────────────
    def _summary_row(label: str, value: str):
        if not value:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        label_run = p.add_run(f"{label}: ")
        label_run.font.size = Pt(10)
        val_run = p.add_run(value)
        val_run.bold = True
        val_run.font.size = Pt(10)

    _summary_row("Brief project description", scope_summary.get("project_description", ""))
    _summary_row("Key objectives",            scope_summary.get("key_objectives", ""))
    _summary_row("Major deliverables",        scope_summary.get("major_deliverables", ""))
    _summary_row("Exclusions (if any)",       scope_summary.get("exclusions", ""))
    doc.add_paragraph()

    # ── Index lookups ─────────────────────────────────────────────────────────
    # Services grouped by their group ID
    svc_by_group: Dict[Any, List[Dict[str, Any]]] = {}
    ungrouped_svcs: List[Dict[str, Any]] = []
    for svc in services:
        gid = svc.get("OpportunityServiceGroupID")
        if gid:
            svc_by_group.setdefault(gid, []).append(svc)
        else:
            ungrouped_svcs.append(svc)

    # Service items (materials) keyed by OpportunityServiceID
    items_by_svc: Dict[Any, List[Dict[str, Any]]] = {}
    for item in service_items:
        sid = (
            item.get("OpportunityServiceID")
            or item.get("ServiceID")
        )
        if sid:
            items_by_svc.setdefault(sid, []).append(item)

    def _service_group_id(grp: Dict[str, Any]) -> Any:
        return (
            grp.get("OpportunityServiceGroupID")
            or grp.get("ServiceGroupID")
            or grp.get("ID")
        )

    def _group_name(grp: Dict[str, Any]) -> str:
        # Confirmed field name from API log
        return grp.get("GroupName") or "—"

    def _group_note(grp: Dict[str, Any]) -> str:
        # Confirmed field name from API log — strip HTML from rich-text field
        return _strip_html(grp.get("GroupDescription"))

    def _item_type(item: Dict[str, Any]) -> str:
        """Normalise ItemType to lowercase — used in both service blocks and section collectors."""
        return (item.get("ItemType") or "").strip().lower()

    def _render_service_block(svc: Dict[str, Any], idx: int):
        """Render one service as a rich block: name, hours, description, materials table."""
        svc_name  = _str(svc.get("DisplayName") or svc.get("ServiceNameAbrOverride"), "—")
        svc_desc  = _strip_html(svc.get("ServiceDescription"))
        op_notes  = _strip_html(svc.get("OperationNotes"))
        hours     = svc.get("ExtendedHours") or svc.get("PerHours")
        svc_id    = svc.get("OpportunityServiceID")
        items     = items_by_svc.get(svc_id, [])

        # Service name bullet
        name_p = doc.add_paragraph(style="List Bullet")
        name_p.paragraph_format.space_before = Pt(6)
        nr = name_p.add_run(svc_name)
        nr.bold = True
        nr.font.size = Pt(10)
        nr.font.color.rgb = _rgb(NAVY)

        if svc_desc:
            dp = doc.add_paragraph()
            dp.paragraph_format.left_indent = Inches(0.3)
            dp.add_run("Description: ").font.size = Pt(9)
            dr = dp.runs[0]
            dp.add_run(svc_desc).font.size = Pt(9)

        if op_notes:
            kp = doc.add_paragraph()
            kp.paragraph_format.left_indent = Inches(0.3)
            klabel = kp.add_run("Key tasks: ")
            klabel.bold = True
            klabel.font.size = Pt(9)
            kr = kp.add_run(op_notes)
            kr.bold = True
            kr.font.size = Pt(9)

        # Hours line
        meta_parts = []
        if hours:
            try:
                meta_parts.append(f"{float(hours):,.1f} Hours")
            except Exception:
                meta_parts.append(str(hours))
        if meta_parts:
            mp = doc.add_paragraph()
            mp.paragraph_format.left_indent = Inches(0.3)
            ml = mp.add_run("Estimated hours: ")
            ml.font.size = Pt(9)
            mv = mp.add_run("  ·  ".join(meta_parts))
            mv.bold = True
            mv.font.size = Pt(9)

        # ── Split items by type ───────────────────────────────────────────────
        mat_items = [i for i in items if _item_type(i) in ("material", "materials", "")]
        # Labor items are intentionally excluded from the document

        def _render_item_table(rows_data: List[Dict[str, Any]]):
            tbl = doc.add_table(rows=1 + len(rows_data), cols=5)
            tbl.style = "Table Grid"
            _section_table_header(tbl, ["Item", "Qty", "Unit", "Vendor", "Notes"])
            col_w = [2.8, 0.65, 1.0, 1.5, 1.55]
            for row in tbl.rows:
                for ci, w in enumerate(col_w):
                    row.cells[ci].width = Inches(w)
            for j, item in enumerate(rows_data, start=1):
                irow = tbl.rows[j]
                bg = GREY if j % 2 == 0 else WHITE
                for cell in irow.cells:
                    _set_cell_bg(cell, bg)
                irow.cells[0].paragraphs[0].add_run(_str(item.get("ItemName"))).font.size          = Pt(9)
                irow.cells[1].paragraphs[0].add_run(str(item.get("ItemQuantity") or "—")).font.size = Pt(9)
                irow.cells[2].paragraphs[0].add_run(_str(item.get("AllocationUnitTypeName"))).font.size = Pt(9)
                irow.cells[3].paragraphs[0].add_run(catalog_vendors.get(item.get("CatalogItemID"), "")).font.size = Pt(9)
                irow.cells[4].paragraphs[0].add_run(_str(_strip_html(item.get("EstimatingNotes") or item.get("ItemDescription")), "")).font.size = Pt(9)

        # Materials table (no labor, no equipment, no subs)
        if mat_items:
            _render_item_table(mat_items)

        doc.add_paragraph()  # spacer after each service block

    # Lists to collect equipment and sub items across all services for Sections 6 & 7
    # Each entry: {"phase": group_name_or_svc_name, "item": item_dict}
    equip_rows: List[Dict[str, Any]] = []
    sub_rows:   List[Dict[str, Any]] = []

    def _collect_typed_items(svc: Dict[str, Any], phase_label: str):
        """Pull equipment and sub items from this service into the cross-section lists."""
        svc_id = svc.get("OpportunityServiceID")
        items  = items_by_svc.get(svc_id, [])
        for item in items:
            t = _item_type(item)
            if t in ("equipment",):
                equip_rows.append({"phase": phase_label, "item": item})
            elif t in ("subcontractor", "sub", "subtrade", "subcontract"):
                sub_rows.append({"phase": phase_label, "item": item})

    if service_groups or ungrouped_svcs:
        svc_counter = [0]

        for grp in service_groups:
            gid      = _service_group_id(grp)
            grp_svcs = svc_by_group.get(gid, [])
            grp_label = _group_name(grp)

            # Group header (navy rule) — name + total hours
            grp_hours = grp.get("ExtendedHours")
            grp_meta  = []
            if grp_hours:
                try:    grp_meta.append(f"{float(grp_hours):,.1f} hrs")
                except: grp_meta.append(str(grp_hours))

            gh = doc.add_paragraph()
            gh.paragraph_format.space_before = Pt(12)
            gh.paragraph_format.space_after  = Pt(2)
            gr = gh.add_run(grp_label)
            gr.bold = True
            gr.font.size = Pt(11)
            gr.font.color.rgb = _rgb(NAVY)
            if grp_meta:
                gm = gh.add_run(f"   {' · '.join(grp_meta)}")
                gm.font.size = Pt(9)
                gm.font.color.rgb = _rgb("475569")
            # Underline rule
            pPr = gh._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),   "single")
            bottom.set(qn("w:sz"),    "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), NAVY)
            pBdr.append(bottom)
            pPr.append(pBdr)

            # Group note
            gnote = _group_note(grp)
            if gnote:
                gn = doc.add_paragraph(gnote)
                gn.runs[0].font.size = Pt(9)
                gn.runs[0].italic = True

            for svc in grp_svcs:
                svc_counter[0] += 1
                _render_service_block(svc, svc_counter[0])
                _collect_typed_items(svc, grp_label)

        # Ungrouped services
        for svc in ungrouped_svcs:
            svc_counter[0] += 1
            svc_label = _str(svc.get("DisplayName") or svc.get("ServiceNameAbrOverride"), "Service")
            _render_service_block(svc, svc_counter[0])
            _collect_typed_items(svc, svc_label)

    else:
        doc.add_paragraph("No service lines found in Aspire for this opportunity.").runs[0].font.size = Pt(10)

    # ── Section 3: Project Schedule ────────────────────────────────────────────
    doc.add_page_break()
    _heading("3. Project Schedule")

    # Build ticket → group name lookup:
    # ticket.OpportunityServiceID → service.OpportunityServiceGroupID → group.GroupName
    svc_id_to_group: Dict[Any, str] = {}
    svc_id_to_hours: Dict[Any, float] = {}   # estimated hours from service record
    grp_id_to_name:  Dict[Any, str] = {
        _service_group_id(g): _group_name(g)
        for g in service_groups
    }
    for svc in services:
        sid  = svc.get("OpportunityServiceID")
        gid  = svc.get("OpportunityServiceGroupID")
        if sid:
            if gid and gid in grp_id_to_name:
                svc_id_to_group[sid] = grp_id_to_name[gid]
            hrs = svc.get("ExtendedHours") or svc.get("PerHours")
            if hrs is not None:
                try:
                    svc_id_to_hours[sid] = float(hrs)
                except Exception:
                    pass

    if tickets:
        # +1 header row + 1 totals row
        sched_table = doc.add_table(rows=1 + len(tickets) + 1, cols=6)
        sched_table.style = "Table Grid"
        _section_table_header(sched_table, [
            "Service Group / Phase", "Ticket #", "Scheduled Date", "Est. Hours", "Actual Hours", "Status"
        ])
        # Column widths
        widths = [2.3, 0.7, 1.1, 0.9, 0.9, 0.6]
        for ci, w in enumerate(widths):
            for row in sched_table.rows:
                row.cells[ci].width = Inches(w)

        total_est = 0.0
        total_act = 0.0
        has_est   = False
        has_act   = False

        for i, tk in enumerate(tickets, start=1):
            row   = sched_table.rows[i]
            status = _str(tk.get("WorkTicketStatusName"), "")
            if "complete" in status.lower():
                bg = GREEN
            elif status.lower() in ("scheduled", "in progress"):
                bg = AMBER
            else:
                bg = WHITE
            for cell in row.cells:
                _set_cell_bg(cell, bg)

            # Build Phase / Title: prefer group name, append ticket title if present
            tk_svc_id  = tk.get("OpportunityServiceID")
            grp_label  = svc_id_to_group.get(tk_svc_id, "")
            tk_title   = _str(tk.get("WorkTicketTitle") or tk.get("WorkTicketName"), "")
            if grp_label and tk_title and tk_title != "—":
                phase_label = f"{grp_label}  —  {tk_title}"
            elif grp_label:
                phase_label = grp_label
            else:
                phase_label = tk_title

            cells = row.cells
            cells[0].paragraphs[0].add_run(phase_label).font.size = Pt(9)
            cells[1].paragraphs[0].add_run(
                _str(tk.get("WorkTicketNumber"))
            ).font.size = Pt(9)
            cells[2].paragraphs[0].add_run(
                _fmt_date(tk.get("ScheduledStartDate") or tk.get("ScheduledDate"))
            ).font.size = Pt(9)

            # Estimated hours — try ticket-level fields first, fall back to service record
            est_hrs_raw = (
                tk.get("EstimatedLaborHours")
                or tk.get("EstimatedHours")
                or tk.get("BudgetedHours")
                or tk.get("ScheduledManHours")
                or tk.get("ScheduledHours")
                or tk.get("ManHours")
                or tk.get("TotalHours")
            )
            try:
                est_hrs = float(est_hrs_raw) if est_hrs_raw is not None else None
            except Exception:
                est_hrs = None
            if est_hrs is None and tk_svc_id:
                est_hrs = svc_id_to_hours.get(tk_svc_id)
            if est_hrs is not None:
                total_est += est_hrs
                has_est = True
            est_str = f"{est_hrs:,.1f}" if est_hrs is not None else "—"
            cells[3].paragraphs[0].add_run(est_str).font.size = Pt(9)

            # Actual hours — try all known Aspire field name variants
            act_hrs_raw = (
                tk.get("ActualLaborHours")
                or tk.get("ActualHours")
                or tk.get("CompletedHours")
                or tk.get("TotalActualHours")
                or tk.get("ActualManHours")
                or tk.get("InvoicedHours")
            )
            try:
                act_hrs = float(act_hrs_raw) if act_hrs_raw is not None else None
            except Exception:
                act_hrs = None
            if act_hrs is not None:
                total_act += act_hrs
                has_act = True
            act_str = f"{act_hrs:,.1f}" if act_hrs is not None else "—"
            cells[4].paragraphs[0].add_run(act_str).font.size = Pt(9)
            cells[5].paragraphs[0].add_run(status).font.size = Pt(9)

        # ── Totals row ────────────────────────────────────────────────────────
        tot_row = sched_table.rows[len(tickets) + 1]
        for cell in tot_row.cells:
            _set_cell_bg(cell, LIGHT)
        def _tot_run(cell, text: str, bold: bool = True):
            r = cell.paragraphs[0].add_run(text)
            r.bold = bold
            r.font.size = Pt(9)
            r.font.color.rgb = _rgb(NAVY)
        _tot_run(tot_row.cells[0], "TOTAL")
        _tot_run(tot_row.cells[1], "")
        _tot_run(tot_row.cells[2], "")
        _tot_run(tot_row.cells[3], f"{total_est:,.1f}" if has_est else "—")
        _tot_run(tot_row.cells[4], f"{total_act:,.1f}" if has_act else "—")
        _tot_run(tot_row.cells[5], "")

        doc.add_paragraph()
        # Legend
        leg_p = doc.add_paragraph()
        leg_p.add_run("Legend:  ").bold = True
        leg_p.add_run("  Green = Complete   ")
        leg_p.add_run("  Amber = Scheduled / In Progress").font.size = Pt(9)
        leg_p.runs[-1].font.size = Pt(9)
        for r in leg_p.runs:
            r.font.size = Pt(9)
    else:
        doc.add_paragraph("No work tickets found in Aspire for this opportunity.").runs[0].font.size = Pt(10)

    # ── Section 4: Materials & Procurement ───────────────────────────────────
    doc.add_page_break()
    _heading("4. Materials & Procurement")

    if receipts:
        for receipt in receipts:
            vendor    = _str(receipt.get("VendorName") or receipt.get("Vendor"), "Unknown Vendor")
            rcpt_date = _fmt_date(receipt.get("ReceivedDate") or receipt.get("ReceiptDate"))
            total     = _fmt_money(receipt.get("TotalAmount") or receipt.get("Total"))
            rcpt_num  = _str(receipt.get("ReceiptNumber") or receipt.get("ReceiptID"), "")

            sub_h = doc.add_paragraph()
            sr = sub_h.add_run(f"{vendor}  —  Receipt #{rcpt_num}  ({rcpt_date})  Total: {total}")
            sr.bold = True
            sr.font.size = Pt(10)
            sr.font.color.rgb = _rgb(NAVY)

            items_raw = receipt.get("ReceiptItems") or []
            if isinstance(items_raw, dict):
                items_raw = items_raw.get("value", [])

            if items_raw:
                item_table = doc.add_table(rows=1 + len(items_raw), cols=4)
                item_table.style = "Table Grid"
                _section_table_header(item_table, ["Item / Description", "Qty", "Unit Cost", "Total"])
                for j, item in enumerate(items_raw, start=1):
                    irow = item_table.rows[j]
                    bg = GREY if j % 2 == 0 else WHITE
                    for cell in irow.cells:
                        _set_cell_bg(cell, bg)
                    qty      = item.get("Quantity") or item.get("Qty") or "—"
                    unit_c   = _fmt_money(item.get("UnitCost") or item.get("UnitPrice"))
                    line_tot = _fmt_money(item.get("TotalCost") or item.get("LineTotal") or item.get("Amount"))
                    desc     = _str(item.get("Description") or item.get("ItemName") or item.get("Name"), "")
                    irow.cells[0].paragraphs[0].add_run(desc).font.size       = Pt(9)
                    irow.cells[1].paragraphs[0].add_run(str(qty)).font.size   = Pt(9)
                    irow.cells[2].paragraphs[0].add_run(unit_c).font.size     = Pt(9)
                    irow.cells[3].paragraphs[0].add_run(line_tot).font.size   = Pt(9)
            else:
                doc.add_paragraph("  (No line items available)").runs[0].font.size = Pt(9)
            doc.add_paragraph()
    else:
        doc.add_paragraph("No purchase receipts found in Aspire for this opportunity.").runs[0].font.size = Pt(10)

    # ── Section 5: Design & Plans ─────────────────────────────────────────────
    doc.add_page_break()
    _heading("5. Design & Plans")
    design_items = [
        "Site Plan / Drawing",
        "Landscape / Planting Plan",
        "Irrigation Plan",
        "Lighting Plan",
        "Grading / Drainage Plan",
        "Structural / Engineering Drawings",
        "Permits Obtained",
        "Permit Posted on Site",
        "HOA Approval",
        "Client-Approved Revisions",
    ]
    dp_table = doc.add_table(rows=1 + len(design_items), cols=3)
    dp_table.style = "Table Grid"
    _section_table_header(dp_table, ["Document / Item", "Included (Y/N)", "Notes"])
    dp_table.columns[0].width = Inches(3.0)
    dp_table.columns[1].width = Inches(1.2)
    dp_table.columns[2].width = Inches(2.3)

    for i, item in enumerate(design_items, start=1):
        row = dp_table.rows[i]
        bg = GREY if i % 2 == 0 else WHITE
        for cell in row.cells:
            _set_cell_bg(cell, bg)
        row.cells[0].paragraphs[0].add_run(item).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run("").font.size   = Pt(9)
        row.cells[2].paragraphs[0].add_run("").font.size   = Pt(9)
    doc.add_paragraph()

    # ── Section 6: Equipment Requirements ────────────────────────────────────
    _heading("6. Equipment Requirements")

    if equip_rows:
        # Group by phase label
        from itertools import groupby as _groupby
        equip_rows.sort(key=lambda r: r["phase"])
        for phase_label, group_iter in _groupby(equip_rows, key=lambda r: r["phase"]):
            phase_items = list(group_iter)
            _heading(phase_label, level=2)
            eq_t = doc.add_table(rows=1 + len(phase_items), cols=4)
            eq_t.style = "Table Grid"
            _section_table_header(eq_t, ["Equipment", "Qty", "Unit", "Notes"])
            col_w_eq = [3.5, 0.65, 1.0, 2.35]
            for row in eq_t.rows:
                for ci, w in enumerate(col_w_eq):
                    row.cells[ci].width = Inches(w)
            for j, entry in enumerate(phase_items, start=1):
                item = entry["item"]
                bg = GREY if j % 2 == 0 else WHITE
                irow = eq_t.rows[j]
                for cell in irow.cells:
                    _set_cell_bg(cell, bg)
                irow.cells[0].paragraphs[0].add_run(_str(item.get("ItemName"))).font.size = Pt(9)
                irow.cells[1].paragraphs[0].add_run(str(item.get("ItemQuantity") or "—")).font.size = Pt(9)
                irow.cells[2].paragraphs[0].add_run(_str(item.get("AllocationUnitTypeName"))).font.size = Pt(9)
                irow.cells[3].paragraphs[0].add_run(_str(_strip_html(item.get("EstimatingNotes") or item.get("ItemDescription")), "")).font.size = Pt(9)
            doc.add_paragraph()
    else:
        doc.add_paragraph("No equipment items found in Aspire for this opportunity.").runs[0].font.size = Pt(10)
        doc.add_paragraph()

    # ── Section 7: Subcontractors ─────────────────────────────────────────────
    _heading("7. Subcontractors")

    if sub_rows:
        from itertools import groupby as _groupby2
        sub_rows.sort(key=lambda r: r["phase"])
        for phase_label, group_iter in _groupby2(sub_rows, key=lambda r: r["phase"]):
            phase_items = list(group_iter)
            _heading(phase_label, level=2)
            sub_t = doc.add_table(rows=1 + len(phase_items), cols=4)
            sub_t.style = "Table Grid"
            _section_table_header(sub_t, ["Subtrade / Item", "Qty", "Unit", "Notes"])
            col_w_sub = [3.5, 0.65, 1.0, 2.35]
            for row in sub_t.rows:
                for ci, w in enumerate(col_w_sub):
                    row.cells[ci].width = Inches(w)
            for j, entry in enumerate(phase_items, start=1):
                item = entry["item"]
                bg = GREY if j % 2 == 0 else WHITE
                irow = sub_t.rows[j]
                for cell in irow.cells:
                    _set_cell_bg(cell, bg)
                irow.cells[0].paragraphs[0].add_run(_str(item.get("ItemName"))).font.size = Pt(9)
                irow.cells[1].paragraphs[0].add_run(str(item.get("ItemQuantity") or "—")).font.size = Pt(9)
                irow.cells[2].paragraphs[0].add_run(_str(item.get("AllocationUnitTypeName"))).font.size = Pt(9)
                irow.cells[3].paragraphs[0].add_run(_str(_strip_html(item.get("EstimatingNotes") or item.get("ItemDescription")), "")).font.size = Pt(9)
            doc.add_paragraph()
    else:
        doc.add_paragraph("No subcontractor items found in Aspire for this opportunity.").runs[0].font.size = Pt(10)
        doc.add_paragraph()

    # ── Section 8: Notes & Special Instructions ────────────────────────────────
    _heading("8. Notes & Special Instructions")
    notes_p = doc.add_paragraph()
    notes_p.add_run("").font.size = Pt(10)
    # Box for hand-writing notes
    for _ in range(8):
        line_p = doc.add_paragraph("_" * 90)
        line_p.runs[0].font.size = Pt(8)
        line_p.runs[0].font.color.rgb = _rgb(MID)

    # ── Footer ────────────────────────────────────────────────────────────────
    from docx.oxml import OxmlElement as _OE
    from docx.oxml.ns import qn as _qn

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fr = fp.add_run(f"{opp_name}  |  Handoff Package  |  {today_str}  |  Page ")
    fr.font.size = Pt(8)
    fr.font.color.rgb = _rgb("6c757d")

    # Add page number field
    fldChar1 = _OE("w:fldChar"); fldChar1.set(_qn("w:fldCharType"), "begin")
    instrText = _OE("w:instrText"); instrText.text = "PAGE"; instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    fldChar2 = _OE("w:fldChar"); fldChar2.set(_qn("w:fldCharType"), "end")
    page_run = fp.add_run()
    page_run._r.append(fldChar1)
    page_run._r.append(instrText)
    page_run._r.append(fldChar2)
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = _rgb("6c757d")

    # ── Serialise to bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Route ─────────────────────────────────────────────────────────────────────

@router.get("/generate")
async def generate_handoff(
    opportunity_number: int = Query(..., description="Aspire Opportunity display number"),
):
    """
    Generate and stream a .docx Handoff Package for the given opportunity.
    All Aspire data is fetched live at generation time.
    """
    _check_credentials()

    logger.info(f"Generating handoff pack for opportunity #{opportunity_number}")

    # Fetch all data in parallel
    import asyncio
    opp = await _fetch_opportunity(opportunity_number)
    opp_id = opp.get("OpportunityID")
    if not opp_id:
        raise HTTPException(status_code=404, detail="OpportunityID missing from Aspire response")

    services, service_groups, tickets, receipts = await asyncio.gather(
        _fetch_services(opp_id),
        _fetch_service_groups(opp_id),
        _fetch_work_tickets(opp_id),
        _fetch_receipts(opp_id),
    )
    # Service items need service IDs — fetch after services are known
    service_ids = [s["OpportunityServiceID"] for s in services if s.get("OpportunityServiceID")]
    service_items = await _fetch_service_items(opp_id, service_ids)

    # Vendor lookup: CatalogItemID → VendorName
    catalog_ids = [i["CatalogItemID"] for i in service_items if i.get("CatalogItemID")]
    catalog_vendors = await _fetch_catalog_vendors(catalog_ids)

    logger.info(
        f"Handoff #{opportunity_number}: {len(service_groups)} groups, "
        f"{len(services)} services, {len(service_items)} items, "
        f"{len(tickets)} tickets, {len(receipts)} receipts"
    )

    # AI-generated scope summary
    group_names  = [g.get("GroupName", "") for g in service_groups if g.get("GroupName")]
    total_hours  = sum(float(g.get("ExtendedHours") or 0) for g in service_groups)
    scope_summary = await _generate_scope_summary(
        opp_name      = _str(opp.get("OpportunityName"), ""),
        prop_name     = _str(opp.get("PropertyName"), ""),
        proposal_desc = _strip_html(opp.get("ProposalDescription") or opp.get("Description")),
        group_names   = group_names,
        total_hours   = total_hours,
    )

    docx_bytes = _build_docx(opp, services, service_groups, service_items, catalog_vendors, scope_summary, tickets, receipts)

    opp_name_slug = (opp.get("OpportunityName") or f"Opportunity-{opportunity_number}") \
        .replace(" ", "_").replace("/", "-")[:60]
    filename = f"Handoff_{opp_name_slug}_{date.today().isoformat()}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
